"""DOCX extractor — paragraphs + tables."""
from pathlib import Path

from docx import Document

from . import ExtractedDoc


def extract_docx(path: Path) -> ExtractedDoc:
    """Extract an :class:`ExtractedDoc` from a ``.docx`` file on disk.

    Paragraphs and table cell text are collected in document order.
    The title is taken from the first Heading-styled paragraph; if none is
    present, falls back to the file stem.
    """
    d = Document(str(path))
    parts: list[str] = []
    title: str | None = None

    for para in d.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "") if para.style is not None else ""
        if title is None and style_name.startswith("Heading"):
            title = text
        parts.append(text)

    for table in d.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append("\t".join(cells))

    return ExtractedDoc(
        title=title or Path(path).stem,
        content="\n\n".join(parts),
        content_type="docx",
        source_path=str(Path(path).resolve()),
        metadata={},
    )
